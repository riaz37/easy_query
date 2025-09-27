# To run this code you need to install the following dependencies:
# pip install google-genai PyPDF2 python-docx openpyxl python-dotenv

import os
import io
import re
import base64
from pathlib import Path
from typing import List, Dict, Any, Tuple, Optional, NamedTuple
from dataclasses import dataclass
from google import genai
from google.genai import types
import PyPDF2
from docx import Document
import openpyxl
from dotenv import load_dotenv
import cloudconvert
import shutil
import tempfile
# Load environment variables
load_dotenv(override=True)

@dataclass
class ChunkBoundary:
    """Represents the boundary information of a chunk"""
    chunk_index: int
    last_sentences: List[str]
    first_sentences: List[str]
    last_table_rows: List[str]
    first_table_rows: List[str]
    has_incomplete_table_start: bool
    has_incomplete_table_end: bool
    page_numbers: List[int]
    
@dataclass
class MergeDecision:
    """Represents a decision on how to merge chunks"""
    chunk1_index: int
    chunk2_index: int
    merge_type: str  # 'text', 'table', 'mixed', 'separate'
    overlap_detected: bool
    instructions: str
    
class SmartDocumentProcessor:
    """
    Advanced document processor with intelligent merging capabilities
    that minimizes LLM calls by using boundary analysis and smart merging.
    """
    
    def __init__(self, api_key: Optional[str] = None, max_pages_per_chunk: int = 10,
                 boundary_sentences: int = 3, boundary_table_rows: int = 3):
        """
        Initialize the smart document processor.
        
        Args:
            api_key: Google Gemini API key (will use env var if not provided)
            max_pages_per_chunk: Maximum pages per chunk for processing
            boundary_sentences: Number of sentences to extract from boundaries
            boundary_table_rows: Number of table rows to extract from boundaries
        """
        self.api_key = api_key or os.getenv("google_api_key")
        self.model_name = os.getenv("google_gemini_name", "wwww-1.5-pro")
        self.max_pages_per_chunk = max_pages_per_chunk
        self.boundary_sentences = boundary_sentences
        self.boundary_table_rows = boundary_table_rows
        
        if not self.api_key:
            raise ValueError("Google API key not found. Set GOOGLE_API_KEY environment variable.")
            
        self.client = genai.Client(api_key=self.api_key)
        
        # Supported file types and their MIME types
        self.supported_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.xls': 'application/vnd.ms-excel',
            '.txt': 'text/plain',
            '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
        }
        
        # Patterns for detecting tables and page numbers
        self.table_patterns = [
            r'\|.*\|',  # Markdown table
            r'^\s*[\w\s]+\s*\|\s*[\w\s]+',  # Simple table format
            r'^\s*\d+\.\s+.*\t.*',  # Tab-separated with numbers
            r'^\s*[A-Z][a-z]+\s+[A-Z][a-z]+\s+\d+',  # Name format tables
        ]
        
        self.page_number_patterns = [
            r'Page\s+(\d+)',
            r'^\s*(\d+)\s*$',
            r'- (\d+) -',
            r'\[(\d+)\]'
        ]
    
    def _get_mime_type(self, file_path: str) -> str:
        """Get MIME type for file based on extension."""
        extension = Path(file_path).suffix.lower()
        return self.supported_types.get(extension, 'application/octet-stream')
    
    def _extract_page_numbers(self, text: str) -> List[int]:
        """Extract page numbers from text."""
        page_numbers = []
        for pattern in self.page_number_patterns:
            matches = re.findall(pattern, text, re.MULTILINE)
            for match in matches:
                try:
                    page_numbers.append(int(match))
                except ValueError:
                    continue
        return sorted(list(set(page_numbers)))
    
    def _extract_sentences(self, text: str, count: int, from_start: bool = True) -> List[str]:
        """Extract sentences from text."""
        # Simple sentence splitting (can be improved with nltk)
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        if from_start:
            return sentences[:count]
        else:
            return sentences[-count:] if len(sentences) >= count else sentences
    
    def _extract_table_rows(self, text: str, count: int, from_start: bool = True) -> List[str]:
        """Extract table rows from text."""
        table_rows = []
        lines = text.split('\n')
        
        for line in lines:
            line = line.strip()
            if any(re.match(pattern, line) for pattern in self.table_patterns):
                table_rows.append(line)
        
        if not table_rows:
            return []
        
        if from_start:
            return table_rows[:count]
        else:
            return table_rows[-count:] if len(table_rows) >= count else table_rows
    
    def _detect_incomplete_tables(self, text: str) -> Tuple[bool, bool]:
        """Detect if text starts or ends with incomplete tables."""
        lines = text.strip().split('\n')
        
        # Check if starts with table continuation
        starts_with_table = False
        for i, line in enumerate(lines[:5]):  # Check first 5 lines
            if any(re.match(pattern, line.strip()) for pattern in self.table_patterns):
                starts_with_table = i == 0  # Table at very beginning
                break
        
        # Check if ends with incomplete table
        ends_with_table = False
        for i, line in enumerate(reversed(lines[-5:])):  # Check last 5 lines
            if any(re.match(pattern, line.strip()) for pattern in self.table_patterns):
                ends_with_table = i == 0  # Table at very end
                break
        
        return starts_with_table, ends_with_table
    
    def _split_pdf(self, file_path: str) -> List[bytes]:
        """Split PDF into manageable chunks."""
        chunks = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            
            for start_page in range(0, total_pages, self.max_pages_per_chunk):
                end_page = min(start_page + self.max_pages_per_chunk, total_pages)
                
                pdf_writer = PyPDF2.PdfWriter()
                for page_num in range(start_page, end_page):
                    pdf_writer.add_page(pdf_reader.pages[page_num])
                
                chunk_stream = io.BytesIO()
                pdf_writer.write(chunk_stream)
                chunks.append(chunk_stream.getvalue())
                
        return chunks
    
    def _split_docx(self, file_path: str) -> List[bytes]:
        """Split DOCX into chunks based on paragraph count."""
        doc = Document(file_path)
        chunks = []
        
        # Estimate paragraphs per page and calculate chunk size
        paragraphs_per_page = 15
        paragraphs_per_chunk = self.max_pages_per_chunk * paragraphs_per_page
        total_paragraphs = len(doc.paragraphs)
        
        for start_idx in range(0, total_paragraphs, paragraphs_per_chunk):
            end_idx = min(start_idx + paragraphs_per_chunk, total_paragraphs)
            
            # Create new document with selected paragraphs
            new_doc = Document()
            for para_idx in range(start_idx, end_idx):
                paragraph = doc.paragraphs[para_idx]
                new_para = new_doc.add_paragraph()
                new_para.text = paragraph.text
                
                # Preserve basic formatting
                for run in paragraph.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
            
            chunk_stream = io.BytesIO()
            new_doc.save(chunk_stream)
            chunks.append(chunk_stream.getvalue())
            
        return chunks
    
    def _split_excel(self, file_path: str) -> List[bytes]:
        """Split Excel file into chunks."""
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        chunks = []
        
        # If multiple sheets, split by sheets
        if len(workbook.sheetnames) > 1:
            sheets_per_chunk = max(1, self.max_pages_per_chunk // 2)
            
            for i in range(0, len(workbook.sheetnames), sheets_per_chunk):
                chunk_sheets = workbook.sheetnames[i:i + sheets_per_chunk]
                
                new_workbook = openpyxl.Workbook()
                # Remove default sheet
                new_workbook.remove(new_workbook.active)
                
                for sheet_name in chunk_sheets:
                    source_sheet = workbook[sheet_name]
                    target_sheet = new_workbook.create_sheet(title=sheet_name)
                    
                    # Copy all data
                    for row in source_sheet.iter_rows(values_only=True):
                        target_sheet.append(row)
                
                chunk_stream = io.BytesIO()
                new_workbook.save(chunk_stream)
                chunks.append(chunk_stream.getvalue())
        else:
            # Single sheet - split by rows
            sheet = workbook.active
            max_row = sheet.max_row
            rows_per_chunk = self.max_pages_per_chunk * 100  # 100 rows per "page"
            
            for start_row in range(1, max_row + 1, rows_per_chunk):
                end_row = min(start_row + rows_per_chunk, max_row + 1)
                
                new_workbook = openpyxl.Workbook()
                new_sheet = new_workbook.active
                
                # Copy headers if not first chunk
                if start_row > 1:
                    header_row = list(sheet.iter_rows(min_row=1, max_row=1, values_only=True))[0]
                    new_sheet.append(header_row)
                
                # Copy data rows
                for row in sheet.iter_rows(min_row=start_row, max_row=end_row-1, values_only=True):
                    new_sheet.append(row)
                
                chunk_stream = io.BytesIO()
                new_workbook.save(chunk_stream)
                chunks.append(chunk_stream.getvalue())
                
        return chunks
    
    def _split_text(self, file_path: str) -> List[bytes]:
        """Split text file into chunks."""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Estimate characters per page
        chars_per_page = 2500
        chars_per_chunk = self.max_pages_per_chunk * chars_per_page
        
        chunks = []
        for start_idx in range(0, len(content), chars_per_chunk):
            end_idx = min(start_idx + chars_per_chunk, len(content))
            chunk_content = content[start_idx:end_idx]
            chunks.append(chunk_content.encode('utf-8'))
            
        return chunks
    
    def split_document(self, file_path: str) -> List[Tuple[bytes, str]]:
        """Split document into chunks based on file type."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = Path(file_path).suffix.lower()
        mime_type = self._get_mime_type(file_path)
        
        if extension not in self.supported_types:
            raise ValueError(f"Unsupported file type: {extension}")
        
        print(f"Splitting {extension} file: {Path(file_path).name}")
        
        # Split based on file type
        if extension == '.pdf':
            chunks = self._split_pdf(file_path)
        elif extension == '.docx':
            chunks = self._split_docx(file_path)
        elif extension in ['.xlsx', '.xls']:
            chunks = self._split_excel(file_path)
        elif extension == '.txt':
            chunks = self._split_text(file_path)
        else:
            # Fallback: read entire file as single chunk
            with open(file_path, 'rb') as f:
                chunks = [f.read()]
        
        return [(chunk, mime_type) for chunk in chunks]
    
    def extract_text_with_boundaries(self, chunk_data: bytes, mime_type: str, 
                                chunk_index: int, total_chunks: int) -> Tuple[str, ChunkBoundary]:
        """Extract text and boundary information from a document chunk."""
        
        extraction_prompt = f"""
        You are processing document chunk {chunk_index + 1} of {total_chunks}.
        
        EXTRACTION REQUIREMENTS:
        
        1. COMPLETE TEXT EXTRACTION:
        - Extract ALL readable text maintaining original structure
        - Preserve headings, paragraphs, lists, and formatting
        - Include page numbers, headers, and footers
        - Maintain logical reading order
        - Mark every new page with [PAGE_START] and [PAGE_BREAK]
        
        2. SEMANTIC STRUCTURE MARKERS:
        - Mark chapter starts: [CHAPTER_START: Title]
        - Mark chapter ends: [CHAPTER_END: Title]
        - Mark section starts: [SECTION_START: Title]
        - Mark section ends: [SECTION_END: Title]
        - Mark subsection starts: [SUBSECTION_START: Title]
        - Mark subsection ends: [SUBSECTION_END: Title]
        - Mark paragraph boundaries: [PARAGRAPH_START] and [PARAGRAPH_END]. Paragraphs should have at least more than two lines.
        - Mark topic changes: [TOPIC_CHANGE: Brief description]
        - Mark list starts/ends: [LIST_START] and [LIST_END]
        - Mark numbered sections: [NUMBERED_SECTION: X.Y.Z Title]
        
        3. TABLE HANDLING:
        - Identify ALL tables in the document
        - Mark table boundaries with [TABLE_START: Description] and [TABLE_END]
        - Preserve complete table structure including headers
        - Use consistent formatting (markdown or CSV style)
        - For partial tables at boundaries, mark as:
            * [TABLE_INCOMPLETE_START] if table begins at chunk start
            * [TABLE_INCOMPLETE_END] if table ends at chunk end
        
        4. CONTENT TYPE MARKERS:
        - Mark introduction sections: [INTRO_START] and [INTRO_END]
        - Mark conclusion sections: [CONCLUSION_START] and [CONCLUSION_END]
        - Mark abstract/summary: [ABSTRACT_START] and [ABSTRACT_END]
        - Mark references/bibliography: [REFERENCES_START] and [REFERENCES_END]
        - Mark appendix sections: [APPENDIX_START: Title] and [APPENDIX_END]
        - Mark code blocks: [CODE_START: Language] and [CODE_END]
        - Mark quotes: [QUOTE_START] and [QUOTE_END]
        
        5. BOUNDARY MARKERS:
        - Mark the beginning of content with [CHUNK_START]
        - Mark the end of content with [CHUNK_END]
        - Preserve all content between these markers
        
        6. SPECIAL ELEMENTS:
        - Mark figures/images as [FIGURE_START: Description] and [FIGURE_END]
        - Mark footnotes as [FOOTNOTE: content]
        - Mark headers/footers: [HEADER: content] and [FOOTER: content]
        - Preserve page numbers and headers
        
        IMPORTANT: Be generous with semantic markers. Better to have too many than too few.
        Each meaningful unit of content should have clear start/end boundaries.
        
        Extract everything systematically and completely with proper semantic structure.
        """
        
        try:
            # Create content for Gemini API
            contents = [
                types.Content(
                    role="user",
                    parts=[
                        types.Part.from_bytes(
                            mime_type=mime_type,
                            data=chunk_data
                        ),
                        types.Part.from_text(text=extraction_prompt)
                    ]
                )
            ]
            
            # Configure generation
            config = types.GenerateContentConfig(
                temperature=0.1,
                response_mime_type="text/plain"
            )
            
            # Generate content
            response_text = ""
            for chunk in self.client.models.generate_content_stream(
                model=self.model_name,
                contents=contents,
                config=config
            ):
                if chunk.text:
                    response_text += chunk.text
            
            extracted_text = response_text.strip()
            
            # Extract boundary information
            boundary = self._extract_boundary_info(extracted_text, chunk_index)
            
            return extracted_text, boundary
            
        except Exception as e:
            error_msg = f"[ERROR_CHUNK_{chunk_index + 1}]: Failed to extract text - {str(e)}"
            print(error_msg)
            
            # Create empty boundary for error case
            boundary = ChunkBoundary(
                chunk_index=chunk_index,
                last_sentences=[],
                first_sentences=[],
                last_table_rows=[],
                first_table_rows=[],
                has_incomplete_table_start=False,
                has_incomplete_table_end=False,
                page_numbers=[]
            )
            
            return error_msg, boundary
    
    def _extract_boundary_info(self, text: str, chunk_index: int) -> ChunkBoundary:
        """Extract boundary information from extracted text."""
        
        # Extract page numbers
        page_numbers = self._extract_page_numbers(text)
        
        # Extract sentences
        first_sentences = self._extract_sentences(text, self.boundary_sentences, from_start=True)
        last_sentences = self._extract_sentences(text, self.boundary_sentences, from_start=False)
        
        # Extract table rows
        first_table_rows = self._extract_table_rows(text, self.boundary_table_rows, from_start=True)
        last_table_rows = self._extract_table_rows(text, self.boundary_table_rows, from_start=False)
        
        # Detect incomplete tables
        has_incomplete_start = '[TABLE_INCOMPLETE_START]' in text
        has_incomplete_end = '[TABLE_INCOMPLETE_END]' in text
        
        # If markers not found, use pattern detection
        if not has_incomplete_start and not has_incomplete_end:
            has_incomplete_start, has_incomplete_end = self._detect_incomplete_tables(text)
        
        return ChunkBoundary(
            chunk_index=chunk_index,
            last_sentences=last_sentences,
            first_sentences=first_sentences,
            last_table_rows=last_table_rows,
            first_table_rows=first_table_rows,
            has_incomplete_table_start=has_incomplete_start,
            has_incomplete_table_end=has_incomplete_end,
            page_numbers=page_numbers
        )
    
    def _decide_merge_strategy(self, boundaries: List[ChunkBoundary]) -> List[MergeDecision]:
        """Decide how to merge chunks based on boundary analysis."""
        
        if len(boundaries) <= 1:
            return []
        
        # For multiple boundaries, make a single LLM call to analyze all transitions
        merge_prompt = self._create_merge_analysis_prompt(boundaries)
        
        try:
            contents = [
                types.Content(
                    role="user",
                    parts=[types.Part.from_text(text=merge_prompt)]
                )
            ]
            
            config = types.GenerateContentConfig(
                temperature=0.1,
                response_mime_type="text/plain"
            )
            
            response_text = ""
            for chunk in self.client.models.generate_content_stream(
                model=self.model_name,
                contents=contents,
                config=config
            ):
                if chunk.text:
                    response_text += chunk.text
            
            # Parse the response to get merge decisions
            decisions = self._parse_merge_decisions(response_text, boundaries)
            return decisions
            
        except Exception as e:
            print(f"Warning: Merge analysis failed - {str(e)}")
            # Fallback to simple concatenation decisions
            return self._create_fallback_decisions(boundaries)
    
    def _create_merge_analysis_prompt(self, boundaries: List[ChunkBoundary]) -> str:
        """Create a comprehensive prompt for merge analysis."""
        
        prompt = f"""
        You are analyzing {len(boundaries)} document chunks to determine the best merging strategy.
        
        ANALYSIS TASK:
        For each pair of consecutive chunks, determine:
        1. Whether there's content overlap or continuation
        2. Whether tables are split across chunks
        3. The best merge strategy
        
        CHUNK BOUNDARY INFORMATION:
        
        """
        
        for i, boundary in enumerate(boundaries):
            prompt += f"""
        CHUNK {i + 1}:
        - Page numbers: {boundary.page_numbers}
        - Incomplete table at start: {boundary.has_incomplete_table_start}
        - Incomplete table at end: {boundary.has_incomplete_table_end}
        - First sentences: {boundary.first_sentences}
        - Last sentences: {boundary.last_sentences}
        - First table rows: {boundary.first_table_rows}
        - Last table rows: {boundary.last_table_rows}
        
        """
        
        prompt += """
        MERGE DECISIONS NEEDED:
        
        For each pair of consecutive chunks (1→2, 2→3, etc.), provide:
        
        CHUNK_PAIR_X:
        MERGE_TYPE: [text|table|mixed|separate]
        OVERLAP_DETECTED: [yes|no]
        INSTRUCTIONS: [specific instructions for merging]
        
        MERGE_TYPE definitions:
        - text: Simple text continuation, merge with paragraph break
        - table: Table continuation, merge rows without headers
        - mixed: Both text and table elements need merging
        - separate: Keep chunks separate, no merging needed
        
        Provide analysis for all consecutive chunk pairs.
        """
        
        return prompt
    
    def _parse_merge_decisions(self, response_text: str, boundaries: List[ChunkBoundary]) -> List[MergeDecision]:
        """Parse LLM response to extract merge decisions."""
        decisions = []
        
        # Simple parsing - look for CHUNK_PAIR patterns
        pair_pattern = r'CHUNK_PAIR_(\d+):\s*MERGE_TYPE:\s*(\w+)\s*OVERLAP_DETECTED:\s*(\w+)\s*INSTRUCTIONS:\s*([^\n]+)'
        matches = re.findall(pair_pattern, response_text, re.MULTILINE | re.DOTALL)
        
        for match in matches:
            pair_num, merge_type, overlap, instructions = match
            pair_index = int(pair_num) - 1
            
            if pair_index < len(boundaries) - 1:
                decision = MergeDecision(
                    chunk1_index=pair_index,
                    chunk2_index=pair_index + 1,
                    merge_type=merge_type.lower(),
                    overlap_detected=overlap.lower() == 'yes',
                    instructions=instructions.strip()
                )
                decisions.append(decision)
        
        # If parsing failed, create fallback decisions
        if not decisions:
            decisions = self._create_fallback_decisions(boundaries)
        
        return decisions
    
    def _create_fallback_decisions(self, boundaries: List[ChunkBoundary]) -> List[MergeDecision]:
        """Create fallback merge decisions based on simple rules."""
        decisions = []
        
        for i in range(len(boundaries) - 1):
            current = boundaries[i]
            next_boundary = boundaries[i + 1]
            
            # Simple rules for fallback
            if current.has_incomplete_table_end and next_boundary.has_incomplete_table_start:
                merge_type = 'table'
                instructions = 'Merge table rows, remove duplicate headers'
            elif current.page_numbers and next_boundary.page_numbers:
                # Check if pages are consecutive
                if any(p + 1 in next_boundary.page_numbers for p in current.page_numbers):
                    merge_type = 'text'
                    instructions = 'Merge with page break'
                else:
                    merge_type = 'separate'
                    instructions = 'Keep separate, non-consecutive pages'
            else:
                merge_type = 'text'
                instructions = 'Standard text merge'
            
            decision = MergeDecision(
                chunk1_index=i,
                chunk2_index=i + 1,
                merge_type=merge_type,
                overlap_detected=False,
                instructions=instructions
            )
            decisions.append(decision)
        
        return decisions
    
    def _smart_merge_chunks(self, extracted_texts: List[str], 
                          decisions: List[MergeDecision]) -> str:
        """Merge chunks based on decisions without additional LLM calls."""
        
        if not extracted_texts:
            return ""
        
        if len(extracted_texts) == 1:
            return extracted_texts[0]
        
        # Start with the first chunk
        merged_text = extracted_texts[0]
        
        # Apply merge decisions sequentially
        for decision in decisions:
            if decision.chunk1_index == 0:  # First merge
                merged_text = self._apply_merge_decision(
                    merged_text, 
                    extracted_texts[decision.chunk2_index], 
                    decision
                )
            else:
                # For subsequent merges, merge with the growing merged_text
                merged_text = self._apply_merge_decision(
                    merged_text,
                    extracted_texts[decision.chunk2_index],
                    decision
                )
        
        return self._final_cleanup(merged_text)

    # 4. ADD FUNCTION TO EXTRACT SEMANTIC STRUCTURE
    def extract_semantic_structure(self, text: str) -> Dict[str, List[Dict[str, Any]]]:
        """Extract semantic structure from marked text for RAG preparation."""
        
        structure = {
            'chapters': [],
            'sections': [],
            'paragraphs': [],
            'tables': [],
            'figures': [],
            'lists': [],
            'code_blocks': [],
            'quotes': [],
            'references': []
        }
        
        # Extract chapters
        chapter_pattern = r'\[CHAPTER_START: ([^\]]+)\](.*?)\[CHAPTER_END'
        chapters = re.findall(chapter_pattern, text, re.DOTALL)
        for i, (title, content) in enumerate(chapters):
            structure['chapters'].append({
                'id': f'chapter_{i+1}',
                'title': title.strip(),
                'content': content.strip(),
                'start_pos': text.find(f'[CHAPTER_START: {title}]'),
                'word_count': len(content.split())
            })
        
        # Extract sections
        section_pattern = r'\[SECTION_START: ([^\]]+)\](.*?)(?=\[SECTION_END|\[CHAPTER_END|\[SECTION_START|$)'
        sections = re.findall(section_pattern, text, re.DOTALL)
        for i, (title, content) in enumerate(sections):
            structure['sections'].append({
                'id': f'section_{i+1}',
                'title': title.strip(),
                'content': content.strip(),
                'start_pos': text.find(f'[SECTION_START: {title}]'),
                'word_count': len(content.split())
            })
        
        # Extract paragraphs
        paragraph_pattern = r'\[PARAGRAPH_START\](.*?)\[PARAGRAPH_END\]'
        paragraphs = re.findall(paragraph_pattern, text, re.DOTALL)
        for i, content in enumerate(paragraphs):
            if content.strip():
                structure['paragraphs'].append({
                    'id': f'paragraph_{i+1}',
                    'content': content.strip(),
                    'start_pos': text.find(f'[PARAGRAPH_START]{content}'),
                    'word_count': len(content.split())
                })
        
        # Extract tables
        table_pattern = r'\[TABLE_START: ([^\]]*)\](.*?)\[TABLE_END\]'
        tables = re.findall(table_pattern, text, re.DOTALL)
        for i, (description, content) in enumerate(tables):
            structure['tables'].append({
                'id': f'table_{i+1}',
                'description': description.strip(),
                'content': content.strip(),
                'start_pos': text.find(f'[TABLE_START: {description}]'),
                'row_count': len([line for line in content.split('\n') if line.strip()])
            })
        
        return structure


    # 5. ADD METHOD TO PROCESS DOCUMENT WITH SEMANTIC EXTRACTION
    def process_document_with_semantics(self, file_path: str, output_dir: Optional[str] = None) -> Tuple[str, Dict]:
        """Process document and return both text and semantic structure."""
        
        # Process document normally
        extracted_text = self.process_document(file_path, output_dir)
        
        # Extract semantic structure
        semantic_structure = self.extract_semantic_structure(extracted_text)
        
        return extracted_text, semantic_structure

    def _convert_to_pdf_if_needed(self, file_path: str) -> str:
        """Convert DOCX/PPTX to PDF if needed. Returns path to PDF file."""
        extension = Path(file_path).suffix.lower()
        load_dotenv(override=True)
        # Configure CloudConvert API from environment variable
        api_key = os.getenv("converter_api_key")
        if not api_key:
            raise RuntimeError("converter_api_key not found in environment variables")

        cloudconvert.configure(api_key=api_key, sandbox=False)
        # Only convert if it's DOCX or PPTX
        if extension not in ['.docx', '.pptx']:
            return file_path  # Return original file path
        
        # Create output PDF path
        pdf_path = file_path.replace(extension, '.pdf')
        
        # Skip conversion if PDF already exists
        if os.path.exists(pdf_path):
            print(f"Using existing PDF: {pdf_path}")
            return pdf_path
        
        print(f"Converting {extension} to PDF...")
        
        try:
            # Create CloudConvert job
            job = cloudconvert.Job.create(payload={
                "tasks": {
                    'upload-my-file': {'operation': 'import/upload'},
                    'convert-my-file': {
                        'operation': 'convert',
                        'input': 'upload-my-file',
                        'output_format': 'pdf'
                    },
                    'export-my-file': {
                        'operation': 'export/url',
                        'input': 'convert-my-file'
                    }
                }
            })
            
            # Upload and convert
            upload_task_id = job['tasks'][0]['id']
            upload_task = cloudconvert.Task.find(id=upload_task_id)
            cloudconvert.Task.upload(file_name=file_path, task=upload_task)
            
            # Wait and download
            job = cloudconvert.Job.wait(id=job['id'])
            export_task = next(t for t in job['tasks'] if t['name'] == 'export-my-file')
            export_task = cloudconvert.Task.wait(id=export_task['id'])
            file_url = export_task['result']['files'][0]['url']
            
            # Download converted PDF
            cloudconvert.download(filename=pdf_path, url=file_url)
            print(f"✓ Converted to PDF: {pdf_path}")
            return pdf_path
            
        except Exception as e:
            raise Exception(f"PDF conversion failed: {str(e)}")

    def _apply_merge_decision(self, text1: str, text2: str, decision: MergeDecision) -> str:
        """Apply a specific merge decision to two text chunks."""
        
        if decision.merge_type == 'separate':
            return text1 + '\n\n[CHUNK_SEPARATOR]\n\n' + text2
        
        elif decision.merge_type == 'table':
            # Merge table rows
            return self._merge_tables(text1, text2)
        
        elif decision.merge_type == 'mixed':
            # Handle mixed content
            return self._merge_mixed_content(text1, text2)
        
        else:  # text merge
            # Remove chunk markers and merge
            clean_text1 = text1.replace('[CHUNK_END]', '').strip()
            clean_text2 = text2.replace('[CHUNK_START]', '').strip()
            
            if decision.overlap_detected:
                # Try to remove overlap
                return self._merge_with_overlap_removal(clean_text1, clean_text2)
            else:
                return clean_text1 + '\n\n' + clean_text2
    
    def _merge_tables(self, text1: str, text2: str) -> str:
        """Merge tables across chunks."""
        
        # Find table boundaries
        table_end_pattern = r'\[TABLE_END\]'
        table_start_pattern = r'\[TABLE_START\]'
        
        # If text1 ends with incomplete table and text2 starts with incomplete table
        if '[TABLE_INCOMPLETE_END]' in text1 and '[TABLE_INCOMPLETE_START]' in text2:
            # Extract table parts
            table1_part = re.split(r'\[TABLE_INCOMPLETE_END\]', text1)[0]
            table2_part = re.split(r'\[TABLE_INCOMPLETE_START\]', text2)[1]
            
            # Merge table parts
            merged_table = table1_part + '\n' + table2_part
            
            # Get non-table parts
            before_table = re.split(r'\[TABLE_START\]', text1)[0] if '[TABLE_START]' in text1 else ''
            after_table = re.split(r'\[TABLE_END\]', text2)[1] if '[TABLE_END]' in text2 else text2
            
            return before_table + '\n[TABLE_START]\n' + merged_table + '\n[TABLE_END]\n' + after_table
        
        # Fallback to simple merge
        return text1 + '\n\n' + text2
    
    def _merge_mixed_content(self, text1: str, text2: str) -> str:
        """Merge mixed content (text and tables)."""
        # This is a simplified version - can be made more sophisticated
        return self._merge_tables(text1, text2)
    
    def _merge_with_overlap_removal(self, text1: str, text2: str) -> str:
        """Merge texts with overlap removal."""
        
        # Simple overlap detection - check if end of text1 matches start of text2
        text1_end = text1[-200:].strip()  # Last 200 chars
        text2_start = text2[:200].strip()  # First 200 chars
        
        # If there's significant overlap, remove it
        if len(text1_end) > 50 and len(text2_start) > 50:
            if text1_end in text2_start:
                # Remove overlap from text2
                overlap_index = text2.find(text1_end)
                if overlap_index != -1:
                    text2 = text2[overlap_index + len(text1_end):].strip()
        
        return text1 + '\n\n' + text2
    
    def _final_cleanup(self, text: str) -> str:
        """Final cleanup of merged text while preserving semantic markers."""
        
        # Remove ONLY processing markers, KEEP semantic markers
        processing_markers_to_remove = [
            r'\[CHUNK_START\]',
            r'\[CHUNK_END\]',
            r'\[TABLE_INCOMPLETE_START\]',
            r'\[TABLE_INCOMPLETE_END\]',
            r'\[CHUNK_SEPARATOR\]'
        ]
        
        for pattern in processing_markers_to_remove:
            text = re.sub(pattern, '', text)
        
        # Clean up extra whitespace but preserve semantic structure
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        text = text.strip()
        
        # Validate semantic markers (optional)
        self._validate_semantic_markers(text)
        
        return text

    # 3. ADD NEW FUNCTION TO VALIDATE SEMANTIC MARKERS
    def _validate_semantic_markers(self, text: str) -> None:
        """Validate that semantic markers are properly paired."""
        
        marker_pairs = [
            ('CHAPTER_START', 'CHAPTER_END'),
            ('SECTION_START', 'SECTION_END'),
            ('SUBSECTION_START', 'SUBSECTION_END'),
            ('PARAGRAPH_START', 'PARAGRAPH_END'),
            ('TABLE_START', 'TABLE_END'),
            ('LIST_START', 'LIST_END'),
            ('INTRO_START', 'INTRO_END'),
            ('CONCLUSION_START', 'CONCLUSION_END'),
            ('ABSTRACT_START', 'ABSTRACT_END'),
            ('REFERENCES_START', 'REFERENCES_END'),
            ('APPENDIX_START', 'APPENDIX_END'),
            ('CODE_START', 'CODE_END'),
            ('QUOTE_START', 'QUOTE_END'),
            ('FIGURE_START', 'FIGURE_END')
        ]
        
        warnings = []
        
        for start_marker, end_marker in marker_pairs:
            start_count = len(re.findall(f'\\[{start_marker}.*?\\]', text))
            end_count = len(re.findall(f'\\[{end_marker}.*?\\]', text))
            
            if start_count != end_count:
                warnings.append(f"Unmatched {start_marker}/{end_marker}: {start_count} starts, {end_count} ends")
        
        if warnings:
            print("Semantic marker warnings:")
            for warning in warnings:
                print(f"  - {warning}")

    def process_document(self, file_path: str, output_dir: Optional[str] = None) -> str:
        """Process entire document with smart merging."""
        
        print(f"Processing document: {Path(file_path).name}")
        file_path = self._convert_to_pdf_if_needed(file_path)

        # Split document into chunks
        chunks = self.split_document(file_path)
        print(f"Document split into {len(chunks)} chunks")
        
        # Save chunks if output directory specified
        if output_dir:
            self._save_chunks(chunks, file_path, output_dir)
        
        # Extract text and boundaries from each chunk
        extracted_texts = []
        boundaries = []
                # Save file path and chunk count to txt log
        with open("document_chunk_numbers.txt", "a") as f:
            f.write(f"{file_path} |Chunks: {len(chunks)}\n")

        
        for i, (chunk_data, mime_type) in enumerate(chunks):
            print(f"Extracting text from chunk {i+1}/{len(chunks)}...")
            
            extracted_text, boundary = self.extract_text_with_boundaries(
                chunk_data, mime_type, i, len(chunks)
            )
            extracted_texts.append(extracted_text)
            boundaries.append(boundary)
            
            print(f"✓ Chunk {i+1} processed ({len(extracted_text)} characters)")
            
            # NEW: Save extracted text from each chunk as individual txt files
            if output_dir:
                os.makedirs(output_dir, exist_ok=True)
                base_name = Path(file_path).stem
                chunk_text_filename = f"{base_name}_chunk_{i+1:03d}_extracted.txt"
                chunk_text_path = os.path.join(output_dir, chunk_text_filename)
                
                with open(chunk_text_path, 'w', encoding='utf-8') as f:
                    f.write(extracted_text)
                print(f"✓ Saved extracted text: {chunk_text_filename}")
        
        # Analyze boundaries and decide merge strategy (single LLM call)
        print("Analyzing chunk boundaries for optimal merging...")
        merge_decisions = self._decide_merge_strategy(boundaries)
        
        # Smart merge without additional LLM calls
        print("Merging chunks using intelligent boundary analysis...")
        final_text = self._smart_merge_chunks(extracted_texts, merge_decisions)
        
        print(f"✓ Document processing completed! Final text: {len(final_text)} characters")
        return final_text
    
    def _save_chunks(self, chunks: List[Tuple[bytes, str]], 
                    original_file: str, output_dir: str) -> None:
        """Save document chunks to files."""
        os.makedirs(output_dir, exist_ok=True)
        base_name = Path(original_file).stem
        extension = Path(original_file).suffix
        
        for i, (chunk_data, _) in enumerate(chunks):
            chunk_filename = f"{base_name}_chunk_{i+1:03d}{extension}"
            chunk_path = os.path.join(output_dir, chunk_filename)
            
            with open(chunk_path, 'wb') as f:
                f.write(chunk_data)
            print(f"Saved: {chunk_filename}")
    
    def save_extracted_text(self, text: str, output_path: str) -> None:
        """Save extracted text to file."""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"Extracted text saved to: {output_path}")
    
    def analyze_document_structure(self, file_path: str) -> Dict[str, Any]:
        """Analyze document structure without full processing."""
        
        chunks = self.split_document(file_path)
        
        analysis = {
            'file_path': file_path,
            'file_type': Path(file_path).suffix.lower(),
            'total_chunks': len(chunks),
            'estimated_pages': self.max_pages_per_chunk * len(chunks),
            'chunk_sizes': [len(chunk[0]) for chunk in chunks],
            'mime_type': chunks[0][1] if chunks else None
        }
        
        return analysis
    
    def process_with_custom_settings(self, file_path: str, 
                                   max_pages_per_chunk: Optional[int] = None,
                                   boundary_sentences: Optional[int] = None,
                                   boundary_table_rows: Optional[int] = None,
                                   output_dir: Optional[str] = None) -> str:
        """Process document with custom settings."""
        
        # Update settings temporarily
        original_max_pages = self.max_pages_per_chunk
        original_sentences = self.boundary_sentences
        original_table_rows = self.boundary_table_rows
        
        if max_pages_per_chunk is not None:
            self.max_pages_per_chunk = max_pages_per_chunk
        if boundary_sentences is not None:
            self.boundary_sentences = boundary_sentences
        if boundary_table_rows is not None:
            self.boundary_table_rows = boundary_table_rows
        
        try:
            result = self.process_document(file_path, output_dir)
        finally:
            # Restore original settings
            self.max_pages_per_chunk = original_max_pages
            self.boundary_sentences = original_sentences
            self.boundary_table_rows = original_table_rows
        
        return result
    
    def batch_process_documents(self, file_paths: List[str], 
                              output_dir: str = "batch_output") -> Dict[str, str]:
        """Process multiple documents in batch."""
        
        os.makedirs(output_dir, exist_ok=True)
        results = {}
        
        for file_path in file_paths:
            try:
                print(f"\n{'='*60}")
                print(f"Processing: {Path(file_path).name}")
                print(f"{'='*60}")
                
                # Process document
                extracted_text = self.process_document(file_path)
                
                # Save result
                output_filename = f"{Path(file_path).stem}_extracted.txt"
                output_path = os.path.join(output_dir, output_filename)
                self.save_extracted_text(extracted_text, output_path)
                
                results[file_path] = output_path
                
            except Exception as e:
                print(f"Error processing {file_path}: {e}")
                results[file_path] = f"ERROR: {str(e)}"
        
        return results


def main():
    """Enhanced example usage of the SmartDocumentProcessor."""
    
    # Initialize processor with custom settings
    try:
        processor = SmartDocumentProcessor(
            max_pages_per_chunk=30,      # Process 5 pages at a time
            boundary_sentences=3,        # Extract 3 sentences from boundaries
            boundary_table_rows=3        # Extract 3 table rows from boundaries
        )
    except ValueError as e:
        print(f"Error: {e}")
        print("Please set your Google API key in environment variable: GOOGLE_API_KEY")
        return
    
    # File to process - UPDATE THIS PATH
    file_path = r"/Users/nilab/Desktop/projects/Knowladge-Base/Annual_Report_2023.pdf"  # Change this to your file path
    output_dir = r"/Users/nilab/Desktop/projects/Knowladge-Base/Parser/document_chunks"           # Directory to save chunks (optional)
    
    # Check if file exists
    if not os.path.exists(file_path):
        print(f"Error: File not found: {file_path}")
        print("\nPlease update the file_path variable with your document path.")
        
        # Show example of analyzing document structure
        print("\nExample: Analyzing document structure...")
        example_files = [
            "document.pdf",
            "report.docx", 
            "data.xlsx",
            "presentation.pptx"
        ]
        
        for example_file in example_files:
            if os.path.exists(example_file):
                analysis = processor.analyze_document_structure(example_file)
                print(f"\nAnalysis for {example_file}:")
                for key, value in analysis.items():
                    print(f"  {key}: {value}")
        
        return
    
    try:
        # # Method 1: Process with default settings
        # print("Method 1: Processing with default settings...")
        # extracted_text = processor.process_document(file_path, output_dir)
        
        # # Save the result
        # output_file = f"{Path(file_path).stem}_extracted.txt"
        # processor.save_extracted_text(extracted_text, output_file)
        
        # Method 2: Process with custom settings
        print("\nMethod 2: Processing with custom settings...")
        custom_extracted = processor.process_with_custom_settings(
            file_path,
            max_pages_per_chunk=5,    # Smaller chunks
            boundary_sentences=5,      # More boundary sentences
            boundary_table_rows=5,     # More table rows for analysis
            output_dir=output_dir + "_custom"
        )
        
        custom_output_file = f"{Path(file_path).stem}_custom_extracted.txt"
        processor.save_extracted_text(custom_extracted, custom_output_file)
        
        # Method 3: Analyze document structure first
        print("\nMethod 3: Document structure analysis...")
        analysis = processor.analyze_document_structure(file_path)
        print("Document Analysis:")
        for key, value in analysis.items():
            print(f"  {key}: {value}")
        
        # Show final summary
        print(f"\n{'='*60}")
        print("PROCESSING SUMMARY")
        print(f"{'='*60}")
        print(f"Source file: {Path(file_path).name}")
        print(f"Custom output: {custom_output_file}")
        print(f"Total characters (custom): {len(custom_extracted):,}")
        print(f"Chunks processed: {analysis['total_chunks']}")
        
        # # Method 4: Batch processing example
        # if input("\nProcess multiple files? (y/n): ").lower() == 'y':
        #     # Add your file paths here
        #     batch_files = [
        #         file_path,  # Include the current file
        #         # Add more files here
        #     ]
            
        #     print("\nBatch processing...")
        #     batch_results = processor.batch_process_documents(batch_files)
            
        #     print("Batch Results:")
        #     for file_path, result in batch_results.items():
        #         print(f"  {Path(file_path).name} -> {result}")
        
    except Exception as e:
        print(f"Error processing document: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    # Run main demo
    main()
    
    # Uncomment to see advanced features demo
    # demo_advanced_features()